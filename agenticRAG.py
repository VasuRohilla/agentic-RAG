import streamlit as st
from dotenv import load_dotenv
import os
import tempfile
from typing import List

# ----------- LANGCHAIN + PINECONE -----------
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone

from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument

from docx import Document

# ----------------- ENV -----------------

load_dotenv()

PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")

if not PINECONE_API_KEY:
    st.error("Missing Pinecone API key.")
    st.stop()


#  AUTO SESSION NAMESPACE

if "namespace" not in st.session_state:
    st.session_state.namespace = f"session_{os.urandom(4).hex()}"


#  CHAT HISTORY STATE

if "messages" not in st.session_state:
    st.session_state.messages = []


#  INITIALIZE MODELS

llm = ChatOllama(
    model="mistral",
    temperature=0.1,
    streaming=False,
    num_ctx=2048,
    timeout=120
)

embeddings = OllamaEmbeddings(
    model="nomic-embed-text"
)

index = Pinecone(api_key=PINECONE_API_KEY).Index("local-index")

vector_store = PineconeVectorStore(
    embedding=embeddings,
    index=index,
    namespace=st.session_state.namespace
)


#  DELETE CURRENT NAMESPACE CONTENT

def clear_current_namespace():
    try:
        index.delete(delete_all=True, namespace=st.session_state.namespace)

    except Exception as e:
        if "Namespace not found" in str(e):
            return
        st.error(f"Pinecone cleanup error: {e}")


#  DOCUMENT PROCESSING

def process_files(files: List):
    documents = []

    for file in files:
        filename = file.name.lower()

        # -------- PDF --------
        if filename.endswith('.pdf'):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name

            loader = PyPDFLoader(tmp_path)
            pages = loader.load()

            for p in pages:
                documents.append(
                    LCDocument(
                        page_content=p.page_content,
                        metadata={
                            "source": file.name,
                            "page": p.metadata.get("page", 0),
                            "type": "pdf"
                        }
                    )
                )

            os.remove(tmp_path)

        # -------- WORD --------
        elif filename.endswith(('.docx', '.doc')):
            doc = Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])

            documents.append(
                LCDocument(
                    page_content=text,
                    metadata={
                        "source": file.name,
                        "type": "word"
                    }
                )
            )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150,
        separators=["\n\n", "\n•", "\n-", "\n", " "]
    )

    return splitter.split_documents(documents)


def add_docs_vectordb(files: List):
    all_splits = process_files(files)

    vector_store.add_documents(
        documents=all_splits,
        batch_size=100
    )


#  QUERY REWRITE

def rewrite_query(query: str) -> str:
    prompt = f"""
You are a search query optimizer.

Rephrase the user question for semantic search.
Keep same intent but add missing context.

Question: {query}

Rewritten query:
"""
    try:
        return llm.invoke(prompt).content.strip()
    except:
        return query


#  RETRIEVAL

def retrieve_context(query: str):

    retriever = vector_store.as_retriever(
        search_type="mmr",
        search_kwargs={
            "k": 8,
            "fetch_k": 20,
            "lambda_mult": 0.7
        }
    )

    docs = retriever.invoke(query)

    context = "\n\n".join(
        f"Source: {d.metadata.get('source')} | Page: {d.metadata.get('page','-')}\n{d.page_content}"
        for d in docs
    )

    return context


#  MAIN RAG CHAIN

def get_ai_response(query: str):

    try:
        better_query = rewrite_query(query)

        context = retrieve_context(better_query)

        if not context.strip():
            return "I don't find this in the document."

        context = context[:6000]

        # PROMPT

        prompt = f"""
You are an expert DOCUMENT REASONING ASSISTANT.

CORE INSTRUCTIONS
1. Use ONLY the information present in the CONTEXT.
2. Do NOT use general knowledge.
3. Do NOT assume or infer missing details.
4. If the exact answer is not in context → reply EXACTLY:
   I don't find this in the document.

REQUIRED STRUCTURE
- Keep your answers concise but insightful.
- Include short bullet points and new lines when possible.
- Preserve numbers, dates, names exactly as written.

Guardrails
- Never add examples not present in context.
- Never explain outside the scope of text.
- If partially available → answer only the available part.

CONTEXT:
{context}

QUESTION:
{query}

ANSWER (follow all rules above):
"""

        response = llm.invoke(prompt)

        return response.content

    except Exception as e:
        return f"LLM Connection Issue: {str(e)}"


#  STREAMLIT UI

st.title("Offline Doc Q&A Buddy")


if st.button("Clear Chat"):
    st.session_state.messages = []
    st.rerun()


uploaded_files = st.file_uploader(
    "Upload Document",
    type=["pdf", "docx", "doc"],
    accept_multiple_files=True
)


if uploaded_files:
    if "indexed_files" not in st.session_state:
        st.session_state.indexed_files = set()

    new_files = [
        f for f in uploaded_files
        if f.name not in st.session_state.indexed_files
    ]

    if new_files:
        with st.spinner("Clearing previous documents..."):
            clear_current_namespace()

        with st.spinner("Indexing new document..."):
            add_docs_vectordb(new_files)

        st.session_state.indexed_files = set(f.name for f in uploaded_files)

        st.success("New document uploaded. Previous data removed!")


# ----- CHAT DISPLAY -----

for msg in st.session_state.messages:

    if msg["role"] == "user":
        with st.chat_message("user", avatar="🧑‍💼"):
            st.markdown(
                f"<div style='text-align: right'>{msg['content']}</div>",
                unsafe_allow_html=True
            )

    else:
        with st.chat_message("assistant", avatar="🤖"):
            st.markdown(
                f"<div style='text-align: left'>{msg['content']}</div>",
                unsafe_allow_html=True
            )


prompt = st.chat_input("Ask from your document...")

if prompt:

    st.session_state.messages.append({
        "role": "user",
        "content": prompt
    })

    with st.chat_message("user", avatar="🧑‍💼"):
        st.markdown(
        f"<div style='text-align: right'>{prompt}</div>",
        unsafe_allow_html=True
    )


    with st.spinner("Thinking..."):
        response = get_ai_response(prompt)

    st.session_state.messages.append({
        "role": "assistant",
        "content": response
    })

    with st.chat_message("assistant", avatar="🤖"):
        st.markdown(
        f"<div style='text-align: left'>{response}</div>",
        unsafe_allow_html=True
    )
